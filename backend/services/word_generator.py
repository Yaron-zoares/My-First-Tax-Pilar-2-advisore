"""
Word Generator Service
Generates Word reports for financial analysis
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from typing import Dict, List, Optional, Any
from datetime import datetime
import logging
import os

from config.settings import settings

logger = logging.getLogger(__name__)

class WordGenerator:
    """Word generator for financial reports"""
    
    def __init__(self):
        """Initialize Word generator"""
        pass
    
    def generate_financial_report(self, data: Dict[str, Any], report_type: str = "standard") -> str:
        """
        Generate financial report in Word format
        
        Args:
            data: Financial data for report generation
            report_type: Type of report (standard, detailed, summary)
            
        Returns:
            Path to generated Word file
        """
        try:
            # Create output directory if it doesn't exist
            output_dir = settings.REPORTS_DIR / "word"
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"financial_report_{report_type}_{timestamp}.docx"
            filepath = output_dir / filename
            
            # Create Word document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Financial Analysis Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add report metadata
            self._add_metadata(doc, data)
            
            # Add financial summary
            self._add_financial_summary(doc, data)
            
            # Add detailed analysis if requested
            if report_type in ["detailed", "comprehensive"]:
                self._add_detailed_analysis(doc, data)
            
            # Add tax calculations
            if report_type in ["detailed", "tax_focused"]:
                self._add_tax_calculations(doc, data)
            
            # Add recommendations
            if report_type in ["detailed", "comprehensive"]:
                self._add_recommendations(doc, data)
            
            # Save document
            doc.save(str(filepath))
            
            logger.info(f"Word report generated successfully: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Word generation error: {str(e)}")
            raise
    
    def _add_metadata(self, doc: Document, data: Dict[str, Any]):
        """Add report metadata"""
        doc.add_heading('Report Information', level=1)
        
        # Create metadata table
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Add metadata rows
        metadata = [
            ["Report Date:", datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
            ["Report Type:", "Financial Analysis"],
            ["Entity Name:", data.get('entity_name', 'N/A')],
            ["Analysis Period:", data.get('period', 'Annual')]
        ]
        
        for i, (key, value) in enumerate(metadata):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = value
        
        doc.add_paragraph()
    
    def _add_financial_summary(self, doc: Document, data: Dict[str, Any]):
        """Add financial summary section"""
        doc.add_heading('Financial Summary', level=1)
        
        # Create summary table
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Add headers
        headers = ["Metric", "Amount (ILS)", "Percentage"]
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
        
        # Add data
        revenue = data.get('revenue', 0)
        expenses = data.get('expenses', 0)
        net_profit = data.get('net_profit', 0)
        
        summary_data = [
            ["Revenue", f"{revenue:,.0f}", "100%"],
            ["Expenses", f"{expenses:,.0f}", f"{(expenses / revenue * 100) if revenue > 0 else 0:.1f}%"],
            ["Net Profit", f"{net_profit:,.0f}", f"{(net_profit / revenue * 100) if revenue > 0 else 0:.1f}%"]
        ]
        
        for i, (metric, amount, percentage) in enumerate(summary_data, 1):
            table.cell(i, 0).text = metric
            table.cell(i, 1).text = amount
            table.cell(i, 2).text = percentage
        
        doc.add_paragraph()
    
    def _add_detailed_analysis(self, doc: Document, data: Dict[str, Any]):
        """Add detailed analysis section"""
        doc.add_heading('Detailed Analysis', level=1)
        
        # Add analysis details
        if 'details' in data:
            details = data['details']
            
            # Categories analysis
            if 'categories' in details:
                doc.add_heading('Financial Categories', level=2)
                category_table = doc.add_table(rows=len(details['categories']) + 1, cols=2)
                category_table.style = 'Table Grid'
                
                # Add headers
                category_table.cell(0, 0).text = "Category"
                category_table.cell(0, 1).text = "Amount (ILS)"
                
                # Add data
                for i, (category, info) in enumerate(details['categories'].items(), 1):
                    category_table.cell(i, 0).text = category.title()
                    category_table.cell(i, 1).text = f"{info['total']:,.0f}"
                
                doc.add_paragraph()
            
            # Risk assessment
            if 'risks' in details:
                doc.add_heading('Risk Assessment', level=2)
                for risk in details['risks']:
                    doc.add_paragraph(f"• {risk}", style='List Bullet')
                doc.add_paragraph()
    
    def _add_tax_calculations(self, doc: Document, data: Dict[str, Any]):
        """Add tax calculations section"""
        doc.add_heading('Tax Calculations', level=1)
        
        # Create tax table
        tax_table = doc.add_table(rows=4, cols=2)
        tax_table.style = 'Table Grid'
        
        # Add headers
        tax_table.cell(0, 0).text = "Tax Metric"
        tax_table.cell(0, 1).text = "Value"
        
        # Add tax data
        tax_data = [
            ["Effective Tax Rate", f"{data.get('effective_tax_rate', 0):.1%}"],
            ["Tax Liability", f"{data.get('tax_liability', 0):,.0f} ILS"],
            ["Tax Efficiency Score", f"{data.get('tax_efficiency_score', 0):.1%}"]
        ]
        
        for i, (metric, value) in enumerate(tax_data, 1):
            tax_table.cell(i, 0).text = metric
            tax_table.cell(i, 1).text = value
        
        doc.add_paragraph()
    
    def _add_recommendations(self, doc: Document, data: Dict[str, Any]):
        """Add recommendations section"""
        doc.add_heading('Recommendations', level=1)
        
        # Add recommendations
        recommendations = data.get('recommendations', [])
        if recommendations:
            for i, rec in enumerate(recommendations, 1):
                doc.add_paragraph(f"{i}. {rec}", style='List Number')
        else:
            doc.add_paragraph("No specific recommendations at this time.")
        
        doc.add_paragraph()
    
    def generate_sample_data(self) -> Dict[str, Any]:
        """
        Generate sample data for testing
        
        Returns:
            Dictionary with sample financial data
        """
        return {
            'entity_name': 'Sample Corporation Ltd.',
            'period': 'Annual 2023',
            'revenue': 1000000,
            'expenses': 750000,
            'net_profit': 250000,
            'effective_tax_rate': 0.23,
            'tax_liability': 57500,
            'tax_efficiency_score': 0.85,
            'details': {
                'categories': {
                    'revenue': {'total': 1000000},
                    'expenses': {'total': 750000},
                    'assets': {'total': 2000000}
                },
                'risks': [
                    'Low profit margin detected',
                    'Revenue concentration risk'
                ]
            },
            'recommendations': [
                'Consider cost optimization strategies',
                'Review tax planning strategies',
                'Address identified financial risks'
            ]
        }
